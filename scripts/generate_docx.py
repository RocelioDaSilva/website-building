from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os

# Paths
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
html_path = os.path.join(ROOT, 'corhed', 'index.html')
output_path = os.path.join(ROOT, 'corhed', 'Corhed_Landing_Edit.docx')

# Try to locate logo referenced in the HTML (relative paths)
possible_logo_paths = [
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.png'),
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.jpg'),
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.svg'),
    os.path.join(ROOT, 'corhed', 'corhed-logo.png'),
]
logo_path = None
for p in possible_logo_paths:
    if os.path.exists(p):
        logo_path = p
        break

"""
generate_docx.py
Generate a Word (.docx) file that contains all visible text from corhed/index.html
including link text, hrefs, form placeholders, input labels, select options and image alt text.
This produces a single editable document the client can use to change any copy.
"""

import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# Paths
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
HTML_PATH = os.path.join(ROOT, 'corhed', 'index.html')
OUTPUT_PATH = os.path.join(ROOT, 'corhed', 'Corhed_Landing_FullText_Edit.docx')

# Candidate logo paths (will embed first found image)
LOGO_CANDIDATES = [
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.png'),
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.jpg'),
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.jpeg'),
    os.path.join(ROOT, 'assets', 'img', 'corhed-logo.svg'),
    os.path.join(ROOT, 'corhed', 'corhed-logo.png'),
]

def find_logo():
    for p in LOGO_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


if not os.path.exists(HTML_PATH):
    raise SystemExit(f"HTML file not found: {HTML_PATH}")

with open(HTML_PATH, 'r', encoding='utf-8') as fh:
    soup = BeautifulSoup(fh, 'lxml')

body = soup.body or soup

# Create Word document
doc = Document()
doc.add_heading('Editable Copy — corhed/index.html', level=1)

# Embed logo if available
logo = find_logo()
if logo:
    try:
        doc.add_picture(logo, width=Inches(1.5))
    except Exception:
        # skip if format not supported (e.g. svg)
        pass

# Title & meta
title = soup.title.string.strip() if soup.title and soup.title.string else ''
meta_desc = soup.find('meta', attrs={'name': 'description'})
desc = meta_desc['content'].strip() if meta_desc and meta_desc.get('content') else ''
if title or desc:
    doc.add_heading('Page Title & Meta', level=2)
    if title:
        p = doc.add_paragraph()
        r = p.add_run('Title: ')
        r.bold = True
        p.add_run(title)
    if desc:
        p = doc.add_paragraph()
        r = p.add_run('Meta description: ')
        r.bold = True
        p.add_run(desc)

# 1) Full visible page text in document order
doc.add_heading('Full Page Visible Text (in document order)', level=2)
full_text = body.get_text(separator='\n', strip=True)
for line in full_text.splitlines():
    line = line.strip()
    if line:
        doc.add_paragraph(line)

# 2) Links (text -> href)
doc.add_heading('Links (text -> href)', level=2)
for a in soup.find_all('a'):
    text = a.get_text(strip=True) or '(no text)'
    href = a.get('href', '')
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(f"{text} → {href}")

# 3) Buttons
buttons = soup.find_all('button')
if buttons:
    doc.add_heading('Buttons (text)', level=2)
    for b in buttons:
        txt = b.get_text(strip=True) or '(no text)'
        doc.add_paragraph(txt, style='List Bullet')

# 4) Form inputs, textareas, selects
doc.add_heading('Form Fields (inputs, textareas, selects)', level=2)

def find_label_for(tag):
    # try label[for=id]
    id_ = tag.get('id')
    if id_:
        lab = soup.find('label', attrs={'for': id_})
        if lab:
            return lab.get_text(strip=True)
    # try parent label
    parent = tag.parent
    if parent and parent.name == 'label':
        return parent.get_text(strip=True)
    return ''

for inp in soup.find_all('input'):
    typ = inp.get('type', 'text')
    name = inp.get('name', '')
    id_ = inp.get('id', '')
    placeholder = inp.get('placeholder', '')
    value = inp.get('value', '')
    label = find_label_for(inp)
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(f"Input [{typ}] name='{name}' id='{id_}' label='{label}' placeholder='{placeholder}' value='{value}'")

for ta in soup.find_all('textarea'):
    name = ta.get('name', '')
    id_ = ta.get('id', '')
    placeholder = ta.get('placeholder', '')
    label = find_label_for(ta)
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(f"Textarea name='{name}' id='{id_}' label='{label}' placeholder='{placeholder}'")

for sel in soup.find_all('select'):
    name = sel.get('name', '')
    id_ = sel.get('id', '')
    label = find_label_for(sel)
    doc.add_paragraph(f"Select name='{name}' id='{id_}' label='{label}'", style='List Bullet')
    for opt in sel.find_all('option'):
        doc.add_paragraph(opt.get_text(strip=True), style='List Bullet')

# 5) Images
imgs = soup.find_all('img')
if imgs:
    doc.add_heading('Images (src → alt)', level=2)
    for img in imgs:
        src = img.get('src', '')
        alt = img.get('alt', '')
        doc.add_paragraph(f"{src} → alt: {alt}", style='List Bullet')

# 6) Convenience editable blocks at the end for quick client editing
doc.add_page_break()
doc.add_heading('Editable Blocks (convenience)', level=1)
doc.add_paragraph('Hero headline: ________________________')
doc.add_paragraph('Hero subheading: ________________________')
doc.add_paragraph('Primary CTA text: ________________________')
doc.add_paragraph('Secondary CTA text: ________________________')
doc.add_paragraph('Services (titles & descriptions):')
for i in range(1, 7):
    doc.add_paragraph(f"{i}. Title: _____________________ Description: ______________________", style='List Bullet')

# Save output
doc.save(OUTPUT_PATH)
print('Saved:', OUTPUT_PATH)
